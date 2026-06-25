"""
Agente Redator ETP — Camada 3 (Arquitetura Modular)

Consolida os dados de Scout, Analista Mercado, Jurisprudência TCU,
Especialista 14.133 e Especialista Técnico num Estudo Técnico
Preliminar robusto e extenso, processando cada secção do art. 18, §1º da
Lei 14.133/2021 de forma isolada para evitar sobrecarga de tokens.
"""

import json
import logging
import os
from datetime import datetime
from docx import Document
from utils.agent_executor import AgentExecutor

logger = logging.getLogger(__name__)


class RedatorETP:

    SECOES_CONFIG = {
        "objeto_etp": "Definição clara, concisa e precisa do objeto da contratação (ex: Sistema de contagem de pessoas por Inteligência Artificial).",
        "i_descricao_necessidade": "Descrição detalhada da necessidade da contratação, demonstrando o problema a ser resolvido, causas, impactos e a real necessidade pública de forma exaustiva.",
        "ii_previsao_pca": "Demonstração do alinhamento da contratação com o Plano de Contratações Anual (PCA) e documentos de planeamento estratégico do órgão.",
        "iii_requisitos_contratacao": "Especificação exaustiva dos requisitos necessários ao atendimento da necessidade, incluindo padrões mínimos de qualidade, SLAs, critérios de eficiência e eficácia.",
        "iv_levantamento_mercado": "Análise técnica detalhada das alternativas possíveis de mercado, contemplando vantagens e desvantagens de cada modelo tecnológico disponível.",
        "v_estimativa_quantidades": "Memória de cálculo detalhada e fundamentada que justifique as quantidades a serem contratadas, acompanhada de métricas utilizadas.",
        "vi_estimativa_valor": "Metodologia detalhada utilizada para a estimativa de custos e o valor total preliminar estimado para a contratação.",
        "vii_descricao_solucoes_existentes": "Descrição técnica completa da solução como um todo, englobando arquitetura, integração e engenharia da melhor alternativa identificada.",
        "viii_justificativa_solucao_escolhida": "Justificativa clara e comparativa técnica e económica dos motivos que levaram à escolha da solução selecionada frente às demais alternativas.",
        "ix_estimativa_impacto_ambiental": "Análise aprofundada dos impactos ambientais diretos e indiretos, elencando medidas mitigadoras e critérios de sustentabilidade aplicáveis.",
        "x_providencias_previas": "Listagem e detalhamento de todas as providências a serem adotadas pela administração antes da celebração do contrato (ex: adequações normativas, vistorias).",
        "xi_contratacoes_correlatas": "Identificação de contratações correlatas, interdependentes ou complementares que possam impactar o sucesso da execução deste objeto.",
        "xii_resultados_pretendidos": "Demonstração clara dos resultados pretendidos em termos de eficiência, economicidade, ganho de produtividade e melhoria do serviço público.",
        "xiii_providencias_adequacao_ambiente": "Mapeamento das providências necessárias para a adequação do ambiente do órgão, infraestrutura tecnológica, elétrica ou lógica para receber a solução.",
        "xiv_analise_riscos": "Matriz e gestão de riscos completa da contratação, apontando riscos identificados, probabilidade, impacto e ações preventivas/mitigadoras (Alinhado ao TCU).",
        "posicionamento_conclusivo": "Declaração e posicionamento conclusivo da equipe de planeamento quanto à viabilidade técnica, jurídica e económica da contratação."
    }

    def executar(self, contexto: dict) -> dict:
        etp_dados_finais = {}
        logger.info("A iniciar a geração modular do ETP com injeção de RAG (TCU/PNCP)...")

        objeto_original = (
            contexto.get("objeto_original")
            or contexto.get("objeto_contratacao")
            or "Sistema de contagem de pessoas por Inteligência Artificial"
        )

        # Extrai os dados do RAG coletados previamente para garantir consistência
        dados_tcu = contexto.get("jurisprudencia_tcu") or contexto.get("Jurisprudencia TCU", {})
        dados_mercado_compras = contexto.get("analista_mercado") or contexto.get("Analista Mercado", {})

        for secao_id, orientacao in self.SECOES_CONFIG.items():
            logger.info(f"A processar secção: {secao_id}")
            
            schema_path = f"schemas/secoes/{secao_id}.json"
            
            # Tratamento robusto para evitar o erro de NoneType caso o schema não exista
            
            if not os.path.exists(schema_path):
                logger.warning(f"Schema específico não encontrado para {secao_id}. Utilizando schema global.")
                schema_path = "schemas/redator_etp_schema.json"

            # Montagem do prompt com blocos isolados e direcionados de RAG
            prompt = f"""
            Você é um Consultor Jurídico e Analista Técnico Especialista em Licitações Públicas, com foco na Lei 14.133/2021.
            Redija a secção técnica '{secao_id}' para o ETP do objeto: "{objeto_original}".

            Diretriz da Secção:
            {orientacao}

            =======================================================================
            PRECECENTES E JURISPRUDÊNCIA DO TCU (RAG Mandatório):
            {json.dumps(dados_tcu, ensure_ascii=False, indent=2)}

            PROCESSOS ANÁLOGOS DO COMPRAS.NET / PNCP (Referências de Mercado):
            {json.dumps(dados_mercado_compras, ensure_ascii=False, indent=2)}
            =======================================================================

            REGRAS CRÍTICAS PARA A RESPOSTA:
            1. Você deve OBRIGATORIAMENTE fundamentar esta seção utilizando os acórdãos do TCU e os processos do PNCP/Compras.net fornecidos acima.
            2. Cite os números dos acórdãos ou os padrões de contratação identificados nos portais oficiais para justificar as escolhas técnicas.
            3. Não resuma. Use parágrafos densos e transcreva as boas práticas e restrições apontadas pelo TCU/PNCP relevantes para esta seção.
            """

            try:
                # O Executor agora recebe o caminho correto corrigido do schema
                resultado_secao = AgentExecutor.executar(prompt=prompt, schema_path=schema_path)
                
                if isinstance(resultado_secao, str):
                    try:
                        resultado_secao = json.loads(resultado_secao)
                    except json.JSONDecodeError:
                        pass
                        
                etp_dados_finais[secao_id] = resultado_secao
            except Exception as e:
                logger.error(f"Falha crítica ao gerar a secção {secao_id}: {e}.")
                etp_dados_finais[secao_id] = {
                    "erro": f"Secção não pôde ser detalhada devido a um erro de processamento interno: {str(e)}"
                }

        self.salvar_docx(etp_dados_finais, objeto_original)
        return etp_dados_finais

    def salvar_docx(self, dados: dict, objeto: str):
        """Lê a árvore de dados complexa gerada e monta uma formatação limpa e profissional no Word."""
        doc = Document()
        doc.add_heading("ESTUDO TÉCNICO PRELIMINAR (ETP)", level=0)
        
        p_sub = doc.add_paragraph()
        p_sub.add_run(f"Objeto da Contratação: {objeto}\n").bold = True
        p_sub.add_run(f"Instrução Processual Licitatória — Lei nº 14.133/2021\nData de Geração: {datetime.now():%d/%m/%Y %H:%M:%S}")

        for indice, (secao_id, conteudo) in enumerate(dados.items(), start=1):
            titulo_secao = f"{indice}. {secao_id.replace('_', ' ').upper()}"
            doc.add_heading(titulo_secao, level=2)

            # Converte e renderiza os dados tipados vindos do Schema granular de cada secção
            if isinstance(conteudo, dict):
                for chave, valor in conteudo.items():
                    # Formata listas internas de forma elegante (ex: vantagens, desvantagens, kpis)
                    if isinstance(valor, list):
                        doc.add_paragraph(f"{chave.replace('_', ' ').title()}:").bold = True
                        for item in valor:
                            if isinstance(item, dict):
                                # Trata sub-objetos dentro de listas como matrizes de risco ou alternativas
                                p_item = doc.add_paragraph(style='List Bullet')
                                for k_sub, v_sub in item.items():
                                    p_item.add_run(f"{k_sub.replace('_', ' ').title()}: ").bold = True
                                    p_item.add_run(f"{v_sub} | ")
                            else:
                                doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        # Campos diretos de string longa
                        if chave.lower() in [secao_id.lower(), "descricao_detalhada", "detalhes_solucao", "conclusao_tecnica"]:
                            doc.add_paragraph(str(valor))
                        else:
                            p = doc.add_paragraph()
                            p.add_run(f"{chave.replace('_', ' ').title()}: ").bold = True
                            p.add_run(str(valor))
            else:
                doc.add_paragraph(str(conteudo))

        # Salva o arquivo final com timestamp exato
        nome_arquivo = f"ETP_Modular_{datetime.now():%Y%m%d_%H%M%S}.docx"
        doc.save(nome_arquivo)
        logger.info(f"Documento Word robusto gerado com sucesso em: {os.path.abspath(nome_arquivo)}")